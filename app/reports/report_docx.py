"""Helper module to export ContractReviewState to DOCX format."""

from __future__ import annotations

import io

from docx import Document

from app import config
from ai_service.utils.masker import unmask_review_state
from ai_service.output_schemas.models import ContractReviewState


from typing import Any

def _safe_str(val: Any) -> str:
    if val is None:
        return ""
    if hasattr(val, "value"):
        return str(val.value)
    return str(val)


def _add_metadata_section(doc: Document, state: ContractReviewState) -> None:
    report = state.final_report
    verdict_str = _safe_str(report.verdict).upper() if report else "N/A"
    risk_str = _safe_str(report.overall_risk_level).upper() if report else "N/A"

    doc.add_heading("Contract Review Report", 0)

    p = doc.add_paragraph()
    p.add_run("Verdict: ").bold = True
    p.add_run(f"{verdict_str}\n")
    p.add_run("Overall Risk Level: ").bold = True
    p.add_run(f"{risk_str}\n")
    p.add_run("Contract ID: ").bold = True
    p.add_run(f"{state.contract_id or 'N/A'}\n")
    p.add_run("Perspective: ").bold = True
    p.add_run(f"{state.perspective or 'Neutral'}\n")

    if state.metadata:
        if state.metadata.document_name:
            p.add_run("Document Name: ").bold = True
            p.add_run(f"{state.metadata.document_name}\n")
        if state.metadata.contract_type:
            p.add_run("Contract Type: ").bold = True
            p.add_run(f"{state.metadata.contract_type}\n")
        if state.metadata.governing_law:
            p.add_run("Governing Law: ").bold = True
            p.add_run(f"{state.metadata.governing_law}\n")

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        report.report_summary if (report and report.report_summary) else "No summary available."
    )


def _add_risk_section(doc: Document, state: ContractReviewState) -> None:
    report = state.final_report
    doc.add_heading("Key Risks Summary", level=1)
    if report and report.key_risks:
        for risk in report.key_risks:
            doc.add_paragraph(risk, style="List Bullet")
    else:
        doc.add_paragraph("No specific key risks identified.")

    if state.risk_scoring and state.risk_scoring.issues:
        doc.add_heading("Detailed Risk Analysis", level=1)
        for issue in state.risk_scoring.issues:
            sev = _safe_str(issue.risk_level).upper()
            score_str = f"{issue.risk_score:.2f}" if issue.risk_score is not None else "N/A"

            dual_scores = []
            if issue.customer_risk_score is not None:
                try:
                    dual_scores.append(f"Customer Risk: {float(issue.customer_risk_score):.2f}")
                except (ValueError, TypeError):
                    dual_scores.append(f"Customer Risk: {issue.customer_risk_score}")
            if issue.vendor_risk_score is not None:
                try:
                    dual_scores.append(f"Vendor Risk: {float(issue.vendor_risk_score):.2f}")
                except (ValueError, TypeError):
                    dual_scores.append(f"Vendor Risk: {issue.vendor_risk_score}")
            score_info = f"Score: {score_str}"
            if dual_scores:
                score_info += f" | {', '.join(dual_scores)}"

            doc.add_heading(f"{issue.clause_type} ({sev} - {score_info})", level=2)

            roles = []
            if issue.benefiting_party:
                roles.append(f"Benefiting Party: {issue.benefiting_party}")
            if issue.burdened_party:
                roles.append(f"Burdened Party: {issue.burdened_party}")
            if issue.liability_holder:
                roles.append(f"Liability Holder: {issue.liability_holder}")
            if issue.decision_controller:
                roles.append(f"Decision Controller: {issue.decision_controller}")
            if roles:
                p_roles = doc.add_paragraph()
                p_roles.add_run("Role Mapping: ").bold = True
                p_roles.add_run(", ".join(roles))

            p_issue = doc.add_paragraph()
            p_issue.add_run("Issue: ").bold = True
            p_issue.add_run(issue.issue)

            if issue.rationale:
                p_rat = doc.add_paragraph()
                p_rat.add_run("Rationale: ").bold = True
                p_rat.add_run(issue.rationale)

            if issue.evidence:
                p_ev = doc.add_paragraph()
                p_ev.add_run("Evidence: ").italic = True
                p_ev.add_run(f"\"{', '.join(issue.evidence)}\"").italic = True

            if issue.negotiation_suggestion:
                p_sug = doc.add_paragraph()
                p_sug.add_run("Negotiation Suggestion: ").bold = True
                p_sug.add_run(issue.negotiation_suggestion)


def _add_red_flags_section(doc: Document, state: ContractReviewState) -> None:
    doc.add_heading("Red Flags", level=1)
    if state.red_flag_detection and state.red_flag_detection.red_flags:
        for flag in state.red_flag_detection.red_flags:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{flag.pattern_name} ").bold = True
            p.add_run(f"({_safe_str(flag.severity).upper()}): {flag.description}")

            roles = []
            if flag.benefiting_party:
                roles.append(f"Benefiting: {flag.benefiting_party}")
            if flag.burdened_party:
                roles.append(f"Burdened: {flag.burdened_party}")
            if flag.liability_holder:
                roles.append(f"Liability: {flag.liability_holder}")
            if flag.decision_controller:
                roles.append(f"Decision Controller: {flag.decision_controller}")
            if roles:
                p_roles = doc.add_paragraph()
                p_roles.add_run("    Role Mapping: ").bold = True
                p_roles.add_run(", ".join(roles))

            if flag.evidence:
                p_ev = doc.add_paragraph()
                p_ev.add_run("    Evidence: ").italic = True
                p_ev.add_run(f"\"{', '.join(flag.evidence)}\"").italic = True
            if flag.safer_alternative:
                p_alt = doc.add_paragraph()
                p_alt.add_run("    Suggested mitigation: ").italic = True
                p_alt.add_run(flag.safer_alternative)
    else:
        doc.add_paragraph("No red flags detected.")


def _add_obligations_section(doc: Document, state: ContractReviewState) -> None:
    doc.add_heading("Key Obligations", level=1)
    if state.obligation_finding and state.obligation_finding.obligations:
        for obl in state.obligation_finding.obligations:
            p = doc.add_paragraph(style="List Bullet")
            if obl.party:
                p.add_run(f"{obl.party}: ").bold = True
            p.add_run(obl.obligation)
            if obl.obligation_type:
                p.add_run(f" (Type: {obl.obligation_type})")

            details = []
            if obl.due_date:
                details.append(f"Due: {obl.due_date}")
            if obl.frequency:
                details.append(f"Frequency: {obl.frequency}")
            if obl.condition:
                details.append(f"Condition: {obl.condition}")
            if details:
                p_det = doc.add_paragraph()
                p_det.add_run("    " + ", ".join(details)).italic = True
    else:
        doc.add_paragraph("No specific obligations extracted.")


def _add_priorities_section(doc: Document, state: ContractReviewState) -> None:
    report = state.final_report
    doc.add_heading("Negotiation Priorities", level=1)
    if report and report.negotiation_priorities:
        for p in sorted(report.negotiation_priorities, key=lambda x: x.priority):
            doc.add_heading(f"{p.priority}. {p.title}", level=2)
            doc.add_paragraph(p.reason)
            if p.recommended_action:
                p_rec = doc.add_paragraph()
                p_rec.add_run("Recommended Action: ").bold = True
                p_rec.add_run(p.recommended_action)
            if p.related_clauses:
                p_rel = doc.add_paragraph()
                p_rel.add_run("Related Clauses: ").bold = True
                p_rel.add_run(", ".join(p.related_clauses))
    else:
        doc.add_paragraph("No negotiation priorities listed.")


def _add_missing_clauses_section(doc: Document, state: ContractReviewState) -> None:
    report = state.final_report
    doc.add_heading("Missing Clauses", level=1)
    if report and report.missing_clauses:
        for m in report.missing_clauses:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{m.category}: ").bold = True
            p.add_run(m.reason or "Not found in contract.")
            if m.impact:
                p_imp = doc.add_paragraph()
                p_imp.add_run(f"    Impact: {m.impact}").italic = True
    else:
        doc.add_paragraph("No missing clauses flagged.")


def _add_plain_english_section(doc: Document, state: ContractReviewState) -> None:
    doc.add_heading("Simplified Clauses (Plain English)", level=1)
    if state.plain_english and state.plain_english.clause_summaries:
        for ps in state.plain_english.clause_summaries:
            doc.add_heading(ps.clause_type, level=2)
            p_trans = doc.add_paragraph()
            p_trans.add_run("Plain English: ").bold = True
            p_trans.add_run(ps.plain_english)
            if ps.why_it_matters:
                p_why = doc.add_paragraph()
                p_why.add_run("Why it matters: ").bold = True
                p_why.add_run(ps.why_it_matters)
            if ps.party_burden:
                p_burd = doc.add_paragraph()
                p_burd.add_run("Burden details: ").bold = True
                p_burd.add_run(ps.party_burden)
    else:
        doc.add_paragraph("No simplified clauses available.")


def export_as_docx(state: ContractReviewState) -> bytes:
    """Generate a clean Microsoft Word document (.docx) of the contract review."""
    if config.ENABLE_SENSITIVE_MASKING:
        state = unmask_review_state(state, config.SENSITIVE_KEYWORDS)

    doc = Document()
    _add_metadata_section(doc, state)
    _add_risk_section(doc, state)
    _add_red_flags_section(doc, state)
    _add_obligations_section(doc, state)
    _add_priorities_section(doc, state)
    _add_missing_clauses_section(doc, state)
    _add_plain_english_section(doc, state)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
