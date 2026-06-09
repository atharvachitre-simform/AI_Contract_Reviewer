"""Helper module to export ContractReviewState to Markdown, PDF, and DOCX formats."""

from __future__ import annotations

import io
import re
from datetime import datetime
from typing import Any
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from reportlab.lib.units import inch
from reportlab.pdfgen import canvas as rl_canvas

from ..models.models import ContractReviewState, RiskLevel, ReviewVerdict


def export_as_markdown(state: ContractReviewState) -> str:
    """Generate a clean markdown report of the contract review state."""
    from src import config
    from .mask import unmask_review_state
    if config.ENABLE_SENSITIVE_MASKING:
        state = unmask_review_state(state, config.SENSITIVE_KEYWORDS)

    report = state.final_report
    
    verdict_str = str(report.verdict.value).upper() if report else "N/A"
    risk_str = str(report.overall_risk_level.value).upper() if report else "N/A"
    
    lines = [
        "# Contract Review Report",
        f"**Verdict**: {verdict_str}  ",
        f"**Overall Risk Level**: {risk_str}  ",
        f"**Contract ID**: {state.contract_id or 'N/A'}  ",
        f"**Perspective**: {state.perspective or 'Neutral'}  ",
    ]
    
    if state.metadata:
        if state.metadata.document_name:
            lines.append(f"**Document Name**: {state.metadata.document_name}  ")
        if state.metadata.contract_type:
            lines.append(f"**Contract Type**: {state.metadata.contract_type}  ")
        if state.metadata.governing_law:
            lines.append(f"**Governing Law**: {state.metadata.governing_law}  ")
            
    lines.extend([
        "",
        "## Executive Summary",
        report.report_summary if (report and report.report_summary) else "No summary available.",
        ""
    ])

    # Key Risks
    lines.append("## Key Risks")
    if report and report.key_risks:
        for risk in report.key_risks:
            lines.append(f"- {risk}")
    else:
        lines.append("*No specific key risks identified.*")
    lines.append("")

    # Red Flags
    lines.append("## Red Flags")
    if state.red_flag_detection and state.red_flag_detection.red_flags:
        for flag in state.red_flag_detection.red_flags:
            lines.append(f"- **{flag.pattern_name}** ({str(flag.severity.value).upper()}): {flag.description}")
            if flag.evidence:
                lines.append(f"  *Evidence:* *\"{', '.join(flag.evidence)}\"*")
            if flag.safer_alternative:
                lines.append(f"  *Suggested fix/alternative:* {flag.safer_alternative}")
    else:
        lines.append("*No red flags detected.*")
    lines.append("")

    # Obligations
    lines.append("## Key Obligations")
    if state.obligation_finding and state.obligation_finding.obligations:
        for obl in state.obligation_finding.obligations:
            party_prefix = f"**{obl.party}** is obligated to: " if obl.party else ""
            lines.append(f"- {party_prefix}{obl.obligation}")
            details = []
            if obl.obligation_type:
                details.append(f"Type: {obl.obligation_type}")
            if obl.due_date:
                details.append(f"Due: {obl.due_date}")
            if obl.frequency:
                details.append(f"Frequency: {obl.frequency}")
            if obl.condition:
                details.append(f"Condition: {obl.condition}")
            if details:
                lines.append(f"  *({', '.join(details)})*")
    else:
        lines.append("*No specific obligations extracted.*")
    lines.append("")

    # Negotiation Priorities
    lines.append("## Negotiation Priorities")
    if report and report.negotiation_priorities:
        for p in sorted(report.negotiation_priorities, key=lambda x: x.priority):
            lines.append(f"### {p.priority}. {p.title}")
            lines.append(p.reason)
            if p.recommended_action:
                lines.append(f"*Recommended action:* {p.recommended_action}")
            if p.related_clauses:
                lines.append(f"*Related clauses:* {', '.join(p.related_clauses)}")
            lines.append("")
    else:
        lines.append("*No negotiation priorities listed.*")
        lines.append("")

    # Missing Clauses
    lines.append("## Missing Clauses")
    if report and report.missing_clauses:
        for m in report.missing_clauses:
            lines.append(f"- **{m.category}**: {m.reason or 'Not found in contract.'}")
            if m.impact:
                lines.append(f"  *Potential Impact:* {m.impact}")
    else:
        lines.append("*No missing clauses flagged.*")
    lines.append("")

    # Plain English Summaries
    lines.append("## Simplified Clauses (Plain English)")
    if state.plain_english and state.plain_english.clause_summaries:
        for ps in state.plain_english.clause_summaries:
            lines.append(f"### {ps.clause_type}")
            lines.append(f"**Plain English translation:** {ps.plain_english}")
            if ps.why_it_matters:
                lines.append(f"*Why it matters:* {ps.why_it_matters}")
            if ps.party_burden:
                lines.append(f"*Party burden details:* {ps.party_burden}")
            lines.append("")
    else:
        lines.append("*No simplified clauses available.*")
        lines.append("")

    lines.append("---")
    lines.append("*Report generated by AI Contract Reviewer*")
    return "\n".join(lines)


def make_callout(text: str, title: str, bg_color: str, border_color: str, title_style: ParagraphStyle, body_style: ParagraphStyle) -> Table:
    content = []
    if title:
        content.append(Paragraph(title, title_style))
        content.append(Spacer(1, 4))
    content.append(Paragraph(text, body_style))
    tbl = Table([[content]], colWidths=[500])
    tbl.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), colors.HexColor(bg_color)),
        ('LEFTPADDING', (0,0), (-1,-1), 12),
        ('RIGHTPADDING', (0,0), (-1,-1), 12),
        ('TOPPADDING', (0,0), (-1,-1), 10),
        ('BOTTOMPADDING', (0,0), (-1,-1), 10),
        ('LINELEFT', (0,0), (-1,-1), 4, colors.HexColor(border_color)),
        ('BOX', (0,0), (-1,-1), 0.5, colors.HexColor(bg_color)),
    ]))
    return tbl


class _NumberedCanvas(rl_canvas.Canvas):
    """Custom canvas that draws page numbers and a footer line on every page."""

    def __init__(self, *args, **kwargs):
        rl_canvas.Canvas.__init__(self, *args, **kwargs)
        self._saved_page_states: list[dict] = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self._draw_footer(num_pages)
            rl_canvas.Canvas.showPage(self)
        rl_canvas.Canvas.save(self)

    def _draw_footer(self, page_count: int) -> None:
        self.saveState()
        w, h = letter
        self.setStrokeColor(colors.HexColor("#E2E8F0"))
        self.setLineWidth(0.5)
        self.line(54, 36, w - 54, 36)
        self.setFont("Helvetica", 7.5)
        self.setFillColor(colors.HexColor("#94A3B8"))
        self.drawString(54, 22, "AI Contract Reviewer — Confidential")
        self.drawRightString(
            w - 54, 22,
            f"Page {self._pageNumber} of {page_count}"
        )
        self.restoreState()


def export_as_pdf(state: ContractReviewState) -> bytes:
    """Generate a high-quality, professional PDF report of the contract review."""
    from src import config
    from .mask import unmask_review_state
    if config.ENABLE_SENSITIVE_MASKING:
        state = unmask_review_state(state, config.SENSITIVE_KEYWORDS)

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=letter,
        leftMargin=54,
        rightMargin=54,
        topMargin=54,
        bottomMargin=54
    )

    styles = getSampleStyleSheet()
    
    # Custom styles to maintain premium look
    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=24,
        leading=28,
        textColor=colors.HexColor("#1A365D"),
        spaceAfter=12
    )
    h1_style = ParagraphStyle(
        'SectionHeading',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=13,
        leading=16,
        textColor=colors.HexColor("#0F172A"),
        spaceBefore=14,
        spaceAfter=4,
        keepWithNext=True
    )
    h2_style = ParagraphStyle(
        'SubsectionHeading',
        parent=styles['Heading3'],
        fontName='Helvetica-Bold',
        fontSize=10.5,
        leading=13,
        textColor=colors.HexColor("#1E293B"),
        spaceBefore=10,
        spaceAfter=4,
        keepWithNext=True
    )
    normal_style = ParagraphStyle(
        'BodyTextCustom',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9.0,
        leading=12.5,
        textColor=colors.HexColor("#334155"),
        spaceAfter=4
    )
    normal_bold_style = ParagraphStyle(
        'BodyTextCustomBold',
        parent=normal_style,
        fontName='Helvetica-Bold',
        textColor=colors.HexColor("#0F172A")
    )
    bullet_style = ParagraphStyle(
        'BulletCustom',
        parent=normal_style,
        leftIndent=15,
        firstLineIndent=-10,
        spaceAfter=4
    )

    story: list[Any] = []

    # --- COVER PAGE ---
    doc_name = state.metadata.document_name if (state.metadata and state.metadata.document_name) else "Contract Review Report"
    if "/" in doc_name or "\\" in doc_name:
        doc_name = doc_name.replace("\\", "/").rsplit("/", 1)[-1]
        
    story.append(Spacer(1, 40))
    story.append(Paragraph("AI CONTRACT AUDIT REPORT", ParagraphStyle('CoverPre', fontName='Helvetica-Bold', fontSize=11, leading=13, textColor=colors.HexColor("#3B82F6"), spaceAfter=8)))
    story.append(Paragraph(doc_name, ParagraphStyle('CoverTitle', fontName='Helvetica-Bold', fontSize=28, leading=34, textColor=colors.HexColor("#0F172A"), spaceAfter=15)))
    
    perspective_str = state.perspective or "Neutral"
    contract_id_str = state.contract_id or "N/A"
    contract_type_str = state.metadata.contract_type if (state.metadata and state.metadata.contract_type) else "General Contract"
    date_str = datetime.now().strftime("%B %d, %Y")
    
    story.append(HRFlowable(width="100%", thickness=2, color=colors.HexColor("#3B82F6"), spaceAfter=25, spaceBefore=5))
    
    cover_meta = f"""
    <b>Contract ID:</b> {contract_id_str}<br/>
    <b>Contract Type:</b> {contract_type_str}<br/>
    <b>Audit Perspective:</b> {perspective_str}<br/>
    <b>Governing Law:</b> {state.metadata.governing_law if (state.metadata and state.metadata.governing_law) else "N/A"}<br/>
    <b>Date of Audit:</b> {date_str}<br/>
    """
    story.append(Paragraph(cover_meta, ParagraphStyle('CoverMeta', fontName='Helvetica', fontSize=10, leading=16, textColor=colors.HexColor("#475569"))))
    story.append(Spacer(1, 180))
    story.append(Paragraph("Confidential report generated by <b>AI Contract Reviewer</b>. All rights reserved.", ParagraphStyle('CoverFoot', fontName='Helvetica-Oblique', fontSize=8.5, leading=11, textColor=colors.HexColor("#94A3B8"))))
    
    from reportlab.platypus import PageBreak
    story.append(PageBreak())

    # --- MAIN REPORT PAGE(S) ---
    report = state.final_report
    verdict_str = str(report.verdict.value).upper() if report else "N/A"
    risk_str = str(report.overall_risk_level.value).upper() if report else "N/A"
    
    # Verdict color
    v_color = "#2ecc71"  # green
    if "review" in verdict_str.lower() or "negotiate" in verdict_str.lower():
        v_color = "#e67e22"  # orange
    elif "redraft" in verdict_str.lower() or "reject" in verdict_str.lower() or "fail" in verdict_str.lower():
        v_color = "#e74c3c"  # red
        
    # Risk color
    r_color = "#2ecc71"  # green
    if "medium" in risk_str.lower():
        r_color = "#f1c40f"  # yellow
    elif "high" in risk_str.lower() or "critical" in risk_str.lower() or "red" in risk_str.lower():
        r_color = "#e74c3c"  # red

    # Verdict & Risk KPI Card Blocks (Side-by-Side)
    v_cell_content = [
        Paragraph("VERDICT", ParagraphStyle('VPre', fontName='Helvetica-Bold', fontSize=8, leading=10, textColor=colors.HexColor("#94A3B8"))),
        Spacer(1, 3),
        Paragraph(f"<b>{verdict_str}</b>", ParagraphStyle('VTitle', fontName='Helvetica-Bold', fontSize=16, leading=18, textColor=colors.HexColor(v_color)))
    ]
    r_cell_content = [
        Paragraph("RISK PROFILE", ParagraphStyle('RPre', fontName='Helvetica-Bold', fontSize=8, leading=10, textColor=colors.HexColor("#94A3B8"))),
        Spacer(1, 3),
        Paragraph(f"<b>{risk_str}</b>", ParagraphStyle('RTitle', fontName='Helvetica-Bold', fontSize=16, leading=18, textColor=colors.HexColor(r_color)))
    ]
    
    kpi_table = Table([[v_cell_content, r_cell_content]], colWidths=[248, 248])
    kpi_table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), colors.HexColor("#F8FAFC")),
        ('ALIGN', (0,0), (-1,-1), 'LEFT'),
        ('VALIGN', (0,0), (-1,-1), 'MIDDLE'),
        ('LEFTPADDING', (0,0), (-1,-1), 14),
        ('RIGHTPADDING', (0,0), (-1,-1), 14),
        ('TOPPADDING', (0,0), (-1,-1), 10),
        ('BOTTOMPADDING', (0,0), (-1,-1), 10),
        ('LINELEFT', (0,0), (0,0), 4, colors.HexColor(v_color)),
        ('LINELEFT', (1,0), (1,0), 4, colors.HexColor(r_color)),
        ('BOX', (0,0), (0,0), 0.5, colors.HexColor("#CBD5E0")),
        ('BOX', (1,0), (1,0), 0.5, colors.HexColor("#CBD5E0")),
    ]))
    story.append(kpi_table)
    story.append(Spacer(1, 12))

    def add_section(title_text):
        story.append(Paragraph(title_text, h1_style))
        story.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor("#E2E8F0"), spaceAfter=8, spaceBefore=4))

    # Executive Summary
    add_section("Executive Summary")
    summary_text = report.report_summary if (report and report.report_summary) else "No summary available."
    story.append(Paragraph(summary_text, normal_style))
    story.append(Spacer(1, 10))

    # Clause Extractor Summary Table
    add_section("Extracted Clauses Overview")
    if state.clause_extraction and state.clause_extraction.clauses:
        table_header = [
            Paragraph("<b>#</b>", normal_bold_style),
            Paragraph("<b>Clause Type</b>", normal_bold_style),
            Paragraph("<b>CUAD Category</b>", normal_bold_style),
            Paragraph("<b>Confidence</b>", normal_bold_style),
        ]
        table_data = [table_header]
        for idx, clause in enumerate(state.clause_extraction.clauses[:40], 1):
            conf = getattr(clause, "confidence", None)
            conf_str = f"{conf:.0%}" if isinstance(conf, float) else (str(conf) if conf else "—")
            table_data.append([
                Paragraph(str(idx), normal_style),
                Paragraph(str(getattr(clause, "clause_type", "Unknown") or "—"), normal_style),
                Paragraph(str(getattr(clause, "cuad_category", "") or "—"), normal_style),
                Paragraph(conf_str, normal_style),
            ])
        clause_table = Table(table_data, colWidths=[28, 190, 190, 60])
        clause_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0F172A")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 8.5),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor("#F8FAFC"), colors.white]),
            ("ALIGN", (0, 0), (-1, -1), "LEFT"),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("LEFTPADDING", (0, 0), (-1, -1), 8),
            ("RIGHTPADDING", (0, 0), (-1, -1), 8),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#CBD5E0")),
            ("LINEBELOW", (0, 0), (-1, -2), 0.3, colors.HexColor("#E2E8F0")),
        ]))
        story.append(clause_table)
        if len(state.clause_extraction.clauses) > 40:
            story.append(Paragraph(
                f"<i>… and {len(state.clause_extraction.clauses) - 40} more clauses (see full analysis below)</i>",
                normal_style
            ))
    else:
        story.append(Paragraph("No clauses were extracted.", normal_style))
    story.append(Spacer(1, 10))

    # Key Risks
    add_section("Key Risks Summary")
    if report and report.key_risks:
        for risk in report.key_risks:
            story.append(Paragraph(f"• {risk}", bullet_style))
    else:
        story.append(Paragraph("No specific key risks identified.", normal_style))
    story.append(Spacer(1, 10))

    # Red Flags (Callouts)
    add_section("Detected Red Flags")
    if state.red_flag_detection and state.red_flag_detection.red_flags:
        for flag in state.red_flag_detection.red_flags:
            sev = str(flag.severity.value).upper()
            if sev == "HIGH":
                bg = "#FEF2F2"
                border = "#EF4444"
            elif sev == "MEDIUM":
                bg = "#FFFBEB"
                border = "#F59E0B"
            else:
                bg = "#F8FAFC"
                border = "#64748B"
            
            flag_title = f"<b>{flag.pattern_name}</b> ({sev})"
            flag_body = f"{flag.description}"
            if flag.evidence:
                flag_body += f"<br/><i>Evidence: \"{', '.join(flag.evidence)}\"</i>"
            if flag.safer_alternative:
                flag_body += f"<br/><b>Suggested Mitigation:</b> {flag.safer_alternative}"
                
            story.append(make_callout(flag_body, flag_title, bg, border, normal_bold_style, normal_style))
            story.append(Spacer(1, 6))
    else:
        story.append(Paragraph("No red flags detected.", normal_style))
    story.append(Spacer(1, 10))

    # Obligations (Merged)
    add_section("Key Obligations")
    if state.obligation_finding and state.obligation_finding.obligations:
        for obl in state.obligation_finding.obligations:
            party_prefix = f"<b>{obl.party}</b>: " if obl.party else ""
            type_suffix = f" <i>(Type: {obl.obligation_type})</i>" if obl.obligation_type else ""
            body = f"{party_prefix}{obl.obligation}{type_suffix}"
            
            details = []
            if obl.due_date:
                details.append(f"Due: {obl.due_date}")
            if obl.frequency:
                details.append(f"Frequency: {obl.frequency}")
            if obl.condition:
                details.append(f"Condition: {obl.condition}")
            if details:
                body += f"<br/><i>Details: {', '.join(details)}</i>"
                
            story.append(make_callout(body, "", "#F8FAFC", "#3B82F6", normal_bold_style, normal_style))
            story.append(Spacer(1, 6))
    else:
        story.append(Paragraph("No specific obligations extracted.", normal_style))
    story.append(Spacer(1, 10))

    # Negotiation Priorities
    add_section("Negotiation Priorities")
    if report and report.negotiation_priorities:
        for p in sorted(report.negotiation_priorities, key=lambda x: x.priority):
            body = f"<b>Priority {p.priority}: {p.title}</b><br/>{p.reason}"
            if p.recommended_action:
                body += f"<br/><b>Recommended Action:</b> {p.recommended_action}"
            if p.related_clauses:
                body += f"<br/><b>Related Clauses:</b> {', '.join(p.related_clauses)}"
            story.append(make_callout(body, "", "#EFF6FF", "#2563EB", normal_bold_style, normal_style))
            story.append(Spacer(1, 6))
    else:
        story.append(Paragraph("No negotiation priorities listed.", normal_style))
    story.append(Spacer(1, 10))

    # Missing Clauses
    add_section("Missing Clauses")
    if report and report.missing_clauses:
        for m in report.missing_clauses:
            body = f"<b>{m.category}</b>: {m.reason or 'Not found.'}"
            if m.impact:
                body += f"<br/><i>Impact: {m.impact}</i>"
            story.append(make_callout(body, "", "#FFFBEB", "#D97706", normal_bold_style, normal_style))
            story.append(Spacer(1, 6))
    else:
        story.append(Paragraph("No missing clauses flagged.", normal_style))
    story.append(Spacer(1, 10))

    # Plain English
    add_section("Simplified Clauses (Plain English)")
    if state.plain_english and state.plain_english.clause_summaries:
        for ps in state.plain_english.clause_summaries:
            body = f"<b>{ps.clause_type}</b><br/><b>Translation:</b> {ps.plain_english}"
            if ps.why_it_matters:
                body += f"<br/><b>Why it matters:</b> {ps.why_it_matters}"
            if ps.party_burden:
                body += f"<br/><b>Burden Details:</b> {ps.party_burden}"
            story.append(make_callout(body, "", "#F8FAFC", "#10B981", normal_bold_style, normal_style))
            story.append(Spacer(1, 6))
    else:
        story.append(Paragraph("No simplified clauses available.", normal_style))

    # Build PDF with numbered canvas for page footers
    doc.build(story, canvasmaker=_NumberedCanvas)
    return buf.getvalue()


def export_as_docx(state: ContractReviewState) -> bytes:
    """Generate a clean Microsoft Word document (.docx) of the contract review."""
    from src import config
    from .mask import unmask_review_state
    if config.ENABLE_SENSITIVE_MASKING:
        state = unmask_review_state(state, config.SENSITIVE_KEYWORDS)

    doc = Document()
    
    report = state.final_report
    verdict_str = str(report.verdict.value).upper() if report else "N/A"
    risk_str = str(report.overall_risk_level.value).upper() if report else "N/A"
    
    # Document header
    doc.add_heading("Contract Review Report", 0)
    
    # Metadata block
    p = doc.add_paragraph()
    p.add_run("Verdict: ").bold = True
    p.add_run(f"{verdict_str}\n")
    p.add_run("Overall Risk Level: ").bold = True
    p.add_run(f"{risk_str}\n")
    p.add_run("Contract ID: ").bold = True
    p.add_run(f"{state.contract_id or 'N/A'}\n")
    p.add_run("Perspective: ").bold = True
    p.add_run(f"{state.perspective or 'Neutral'}\n")

    if state.metadata:
        if state.metadata.document_name:
            p.add_run("Document Name: ").bold = True
            p.add_run(f"{state.metadata.document_name}\n")
        if state.metadata.contract_type:
            p.add_run("Contract Type: ").bold = True
            p.add_run(f"{state.metadata.contract_type}\n")
        if state.metadata.governing_law:
            p.add_run("Governing Law: ").bold = True
            p.add_run(f"{state.metadata.governing_law}\n")
            
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(report.report_summary if (report and report.report_summary) else "No summary available.")

    # Key Risks
    doc.add_heading("Key Risks", level=1)
    if report and report.key_risks:
        for risk in report.key_risks:
            doc.add_paragraph(risk, style="List Bullet")
    else:
        doc.add_paragraph("No specific key risks identified.")

    # Red Flags
    doc.add_heading("Red Flags", level=1)
    if state.red_flag_detection and state.red_flag_detection.red_flags:
        for flag in state.red_flag_detection.red_flags:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{flag.pattern_name} ").bold = True
            p.add_run(f"({str(flag.severity.value).upper()}): {flag.description}")
            if flag.evidence:
                p_ev = doc.add_paragraph()
                p_ev.add_run("    Evidence: ").italic = True
                p_ev.add_run(f"\"{', '.join(flag.evidence)}\"").italic = True
            if flag.safer_alternative:
                p_alt = doc.add_paragraph()
                p_alt.add_run("    Suggested mitigation: ").italic = True
                p_alt.add_run(flag.safer_alternative)
    else:
        doc.add_paragraph("No red flags detected.")

    # Obligations
    doc.add_heading("Key Obligations", level=1)
    if state.obligation_finding and state.obligation_finding.obligations:
        for obl in state.obligation_finding.obligations:
            p = doc.add_paragraph(style="List Bullet")
            if obl.party:
                p.add_run(f"{obl.party}: ").bold = True
            p.add_run(obl.obligation)
            if obl.obligation_type:
                p.add_run(f" (Type: {obl.obligation_type})")
            
            details = []
            if obl.due_date:
                details.append(f"Due: {obl.due_date}")
            if obl.frequency:
                details.append(f"Frequency: {obl.frequency}")
            if obl.condition:
                details.append(f"Condition: {obl.condition}")
            if details:
                p_det = doc.add_paragraph()
                p_det.add_run("    " + ", ".join(details)).italic = True
    else:
        doc.add_paragraph("No specific obligations extracted.")

    # Negotiation Priorities
    doc.add_heading("Negotiation Priorities", level=1)
    if report and report.negotiation_priorities:
        for p in sorted(report.negotiation_priorities, key=lambda x: x.priority):
            doc.add_heading(f"{p.priority}. {p.title}", level=2)
            doc.add_paragraph(p.reason)
            if p.recommended_action:
                p_rec = doc.add_paragraph()
                p_rec.add_run("Recommended Action: ").bold = True
                p_rec.add_run(p.recommended_action)
            if p.related_clauses:
                p_rel = doc.add_paragraph()
                p_rel.add_run("Related Clauses: ").bold = True
                p_rel.add_run(", ".join(p.related_clauses))
    else:
        doc.add_paragraph("No negotiation priorities listed.")

    # Missing Clauses
    doc.add_heading("Missing Clauses", level=1)
    if report and report.missing_clauses:
        for m in report.missing_clauses:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(f"{m.category}: ").bold = True
            p.add_run(m.reason or "Not found in contract.")
            if m.impact:
                p_imp = doc.add_paragraph()
                p_imp.add_run(f"    Impact: {m.impact}").italic = True
    else:
        doc.add_paragraph("No missing clauses flagged.")

    # Plain English
    doc.add_heading("Simplified Clauses (Plain English)", level=1)
    if state.plain_english and state.plain_english.clause_summaries:
        for ps in state.plain_english.clause_summaries:
            doc.add_heading(ps.clause_type, level=2)
            p_trans = doc.add_paragraph()
            p_trans.add_run("Plain English: ").bold = True
            p_trans.add_run(ps.plain_english)
            if ps.why_it_matters:
                p_why = doc.add_paragraph()
                p_why.add_run("Why it matters: ").bold = True
                p_why.add_run(ps.why_it_matters)
            if ps.party_burden:
                p_burd = doc.add_paragraph()
                p_burd.add_run("Burden details: ").bold = True
                p_burd.add_run(ps.party_burden)
    else:
        doc.add_paragraph("No simplified clauses available.")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
